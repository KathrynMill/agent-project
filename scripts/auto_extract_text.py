#!/usr/bin/env python3
"""
自动提取Word文档文本内容
使用python-docx库提取文本和图片信息
"""

import os
import sys
import json
import logging
import base64
from pathlib import Path
from datetime import datetime
import zipfile
from io import BytesIO

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)

# 尝试导入python-docx
try:
    from docx import Document
    from PIL import Image
    DOCX_AVAILABLE = True
    logger.info("python-docx 库可用")
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx 库不可用，将使用备用方法")


class DocxTextExtractor:
    """Word文档文本提取器"""

    def __init__(self, docx_path: str):
        self.docx_path = Path(docx_path)
        self.content = {
            "text": "",
            "paragraphs": [],
            "images": [],
            "metadata": {}
        }
        self.success = False

    def extract_content(self):
        """提取文档内容"""
        try:
            if DOCX_AVAILABLE:
                self._extract_with_docx()
            else:
                self._extract_with_zipfile()

            self.success = True
            logger.info(f"成功提取内容: {self.docx_path.name}")

        except Exception as e:
            logger.error(f"提取失败: {e}")
            self.success = False

    def _extract_with_docx(self):
        """使用python-docx提取"""
        doc = Document(self.docx_path)

        # 提取段落
        paragraphs = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text.strip())

        self.content["text"] = "\n".join(paragraphs)
        self.content["paragraphs"] = paragraphs

        # 提取图片
        self._extract_images_from_docx()

    def _extract_with_zipfile(self):
        """使用zipfile提取（备用方法）"""
        with zipfile.ZipFile(self.docx_path, 'r') as zip_file:
            # 读取文档内容
            document_xml = zip_file.read('word/document.xml')

            # 简单的XML文本提取
            import re
            text_content = re.sub(r'<[^>]+>', ' ', document_xml.decode('utf-8'))
            text_content = ' '.join(text_content.split())

            # 提取段落文本（简化版）
            paragraphs = []
            text_blocks = re.findall(r'([^<>\n]{10,100})', text_content)
            for block in text_blocks:
                if block.strip() and len(block) > 20:  # 过滤太短的文本
                    paragraphs.append(block.strip())

            self.content["text"] = "\n".join(paragraphs)
            self.content["paragraphs"] = paragraphs

            # 提取图片信息
            self._extract_images_from_zipfile(zip_file)

    def _extract_images_from_docx(self):
        """从docx提取图片"""
        if not DOCX_AVAILABLE:
            return

        doc = Document(self.docx_path)
        image_count = 0

        for rel in doc.part.rels.iter():
            if "image" in rel.target_ref:
                try:
                    image_part = doc.part.get_part(rel.target_ref)
                    image_data = image_part.blob

                    # 尝试识别图片类型
                    img = Image.open(BytesIO(image_data))
                    width, height = img.size
                    format_name = img.format

                    images_info = {
                        "index": image_count + 1,
                        "filename": f"image_{image_count + 1}.{format_name.lower()}" if format_name else f"image_{image_count + 1}.jpg",
                        "width": width,
                        "height": height,
                        "format": format_name,
                        "size": len(image_data),
                        "base64": base64.b64encode(image_data).decode('utf-8')
                    }

                    self.content["images"].append(images_info)
                    image_count += 1

                except Exception as e:
                    logger.warning(f"处理图片失败: {e}")

    def _extract_images_from_zipfile(self, zip_file):
        """从zipfile提取图片信息"""
        images_info = []

        try:
            for file in zip_file.filelist:
                if file.filename.startswith('word/media/') and \
                   file.filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):

                    image_data = zip_file.read(file.filename)

                    try:
                        img = Image.open(BytesIO(image_data))
                        width, height = img.size
                        format_name = img.format if img.format else "Unknown"
                    except:
                        # 无法识别图片格式
                        width = height = 0
                        format_name = "Unknown"

                    image_info = {
                        "filename": os.path.basename(file.filename),
                        "internal_path": file.filename,
                        "width": width,
                        "height": height,
                        "format": format_name,
                        "size": len(image_data),
                        "base64": base64.b64encode(image_data).decode('utf-8')
                    }

                    images_info.append(image_info)

        except Exception as e:
            logger.warning(f"从zipfile提取图片失败: {e}")

        self.content["images"] = images_info

    def analyze_content(self):
        """分析提取的内容"""
        text = self.content.get("text", "")
        images = self.content.get("images", [])

        # 基础统计
        char_count = len(text)
        word_count = len(text.split())
        paragraph_count = len(self.content.get("paragraphs", []))

        # 角色推断（从文件名）
        filename = self.docx_path.name
        role = "未知"
        if "柯太太" in filename:
            role = "柯太太"
        elif "柯少爷" in filename:
            role = "柯少爷"
        elif "云晴" in filename:
            role = "云晴"
        elif "零四" in filename:
            role = "零四"
        elif "雾晓" in filename:
            role = "雾晓"
        elif "手册" in filename:
            role = "游戏手册"
        elif "线索" in filename:
            role = "线索材料"

        # 内容特征分析
        features = {
            "has_dialogue": "：" in text or "\"" in text,
            "has_timeline": any(word in text for word in ["时间", "点", "分钟", "小时"]),
            "has_clues": any(word in text for word in ["线索", "证据", "秘密"]),
            "has_death": any(word in text for word in ["死亡", "死", "被杀"])
        }

        self.content["metadata"] = {
            "role": role,
            "char_count": char_count,
            "word_count": word_count,
            "paragraph_count": paragraph_count,
            "image_count": len(images),
            "file_size": self.docx_path.stat().st_size,
            "features": features,
            "estimated_duration_hours": max(1, char_count / 3000),  # 简单估算
            "complexity": "medium" if char_count > 2000 else "simple"
        }

    def print_summary(self):
        """打印提取摘要"""
        if not self.success:
            print(f"❌ 提取失败: {self.docx_path}")
            return

        print(f"📄 文件: {self.docx_path.name}")
        print(f"🎭 角色: {self.content['metadata']['role']}")
        print(f"📝 字数: {self.content['metadata']['char_count']}")
        print(f"📖 段落: {self.content['metadata']['paragraph_count']}")
        print(f"🖼️ 图片: {self.content['metadata']['image_count']} 张")

        if self.content['metadata']['features']['has_dialogue']:
            print("💬 包含对话")
        if self.content['metadata']['features']['has_timeline']:
            print("⏰ 包含时间线")
        if self.content['metadata']['features']['has_clues']:
            print("🔍 包含线索")
        if self.content['metadata']['features']['has_death']:
            print("💀 包含死亡情节")

        # 显示前几段内容
        print("\n📖 文本预览 (前3段):")
        for i, paragraph in enumerate(self.content['paragraphs'][:3]):
            print(f"  {i+1}. {paragraph[:100]}{'...' if len(paragraph) > 100 else ''}")


def process_all_files():
    """处理所有docx文件"""
    docx_files = [
        "/home/ubt/桌面/agent-project/01 绷带女人 柯太太_QQ浏览器转格式.docx",
        "/home/ubt/桌面/agent-project/02 年轻男子 柯少爷_QQ浏览器转格式.docx",
        "/home/ubt/桌面/agent-project/03 女仆 云晴_QQ浏览器转格式.docx",
        "/home/ubt/桌面/agent-project/04 胡茬男人 零四_QQ浏览器转格式 (1).docx",
        "/home/ubt/桌面/agent-project/05 洋裙女子 雾晓_QQ浏览器转格式.docx",
        "/home/ubt/桌面/agent-project/手册_QQ浏览器转格式.docx",
        "/home/ubt/桌面/agent-project/线索_QQ浏览器转格式.docx"
    ]

    print("🎭 柯家庄园谋杀案 - 自动文本提取")
    print("="*50)

    extracted_data = {
        "title": "柯家庄园谋杀案",
        "extracted_at": datetime.now().isoformat(),
        "files": {},
        "full_text": "",
        "all_images": [],
        "metadata": {}
    }

    character_scripts = {}
    manual_data = None
    clues_data = None

    # 处理每个文件
    for docx_file in docx_files:
        if not os.path.exists(docx_file):
            print(f"❌ 文件不存在: {docx_file}")
            continue

        print(f"\n📄 处理文件: {os.path.basename(docx_file)}")
        extractor = DocxTextExtractor(docx_file)
        extractor.extract_content()
        extractor.analyze_content()
        extractor.print_summary()

        if extractor.success:
            role = extractor.content['metadata']['role']
            file_data = {
                "text": extractor.content['text'],
                "metadata": extractor.content['metadata'],
                "images": extractor.content['images']
            }

            extracted_data['files'][os.path.basename(docx_file)] = file_data

            # 分类存储
            if role in ['柯太太', '柯少爷', '云晴', '零四', '雾晓']:
                character_scripts[role] = file_data
            elif role == '游戏手册':
                manual_data = file_data
            elif role == '线索材料':
                clues_data = file_data

            # 添加到完整文本
            if extractor.content['text']:
                extracted_data['full_text'] += f"\n\n=== {role} ===\n"
                extracted_data['full_text'] += extractor.content['text']

            # 收集所有图片
            if extractor.content['images']:
                extracted_data['all_images'].extend(extractor.content['images'])

    # 统计信息
    extracted_data['metadata'] = {
        'total_files': len(extracted_data['files']),
        'character_count': len(character_scripts),
        'has_manual': manual_data is not None,
        'has_clues': clues_data is not None,
        'total_images': len(extracted_data['all_images']),
        'full_text_length': len(extracted_data['full_text'])
    }

    # 保存完整提取数据
    output_dir = Path("/home/ubt/桌面/agent-project/data/extracted")
    output_dir.mkdir(parents=True, exist_ok=True)

    # 保存详细数据
    full_output_file = output_dir / "ke_mansion_murder_full.json"
    with open(full_output_file, 'w', encoding='utf-8') as f:
        json.dump(extracted_data, f, ensure_ascii=False, indent=2)

    # 保存分类数据
    extracted_data['character_scripts'] = character_scripts
    extracted_data['game_manual'] = manual_data
    extracted_data['clues'] = clues_data

    simple_output_file = output_dir / "ke_mansion_murder_simple.json"
    with open(simple_output_file, 'w', encoding='utf-8') as f:
        json.dump({
            "title": extracted_data['title'],
            "character_scripts": {role: data['text'] for role, data in character_scripts.items()},
            "game_manual": manual_data['text'] if manual_data else "",
            "clues": clues_data['text'] if clues_data else "",
            "full_text": extracted_data['full_text'],
            "metadata": extracted_data['metadata']
        }, f, ensure_ascii=False, indent=2)

    print("\n" + "="*50)
    print("✅ 提取完成！")
    print("="*50)
    print(f"📊 统计信息:")
    print(f"   总文件数: {extracted_data['metadata']['total_files']}")
    print(f"   角色数量: {extracted_data['metadata']['character_count']}")
    print(f"   手册文件: {'有' if extracted_data['metadata']['has_manual'] else '无'}")
    print(f"   线索文件: {'有' if extracted_data['metadata']['has_clues'] else '无'}")
    print(f"   总图片数: {extracted_data['metadata']['total_images']}")
    print(f"   文本长度: {extracted_data['metadata']['full_text_length']} 字符")

    print(f"\n📁 输出文件:")
    print(f"   完整数据: {full_output_file}")
    print(f"   简化数据: {simple_output_file}")

    return extracted_data


def create_compressed_sample(full_data):
    """创建压缩样本示例"""
    print("\n🎯 创建压缩样本示例...")

    full_text = full_data['full_text']
    if not full_text:
        print("❌ 没有文本内容，无法创建压缩样本")
        return

    # 简单的压缩逻辑：基于关键词分段压缩
    sections = full_text.split('\n=== ')
    compressed_sections = []

    for section in sections:
        if not section.strip():
            continue

        lines = section.split('\n')
        if len(lines) > 10:
            # 保留前几行和最后几行，压缩中间内容
            if len(lines) > 20:
                compressed_section = '\n'.join(lines[:3] + ['...'] + lines[-2:])
            else:
                # 简单压缩
                compressed_section = '\n'.join(lines[:len(lines)//2])
        else:
            compressed_section = section

        compressed_sections.append(compressed_section)

    compressed_text = '\n\n'.join(compressed_sections)

    # 计算压缩比例
    original_length = len(full_text)
    compressed_length = len(compressed_text)
    compression_ratio = compressed_length / original_length if original_length > 0 else 0.6

    # 创建训练样本
    training_sample = {
        "script_id": "ke_murder_compressed_001",
        "title": "柯家庄园谋杀案",
        "original_script": full_text,
        "compressed_script": compressed_text,
        "compression_ratio": compression_ratio,
        "compression_level": "medium" if compression_ratio > 0.6 else "heavy",
        "logic_integrity": 0.8,  # 假设值
        "story_coherence": 0.75,  # 假设值
        "playability_score": 0.82,  # 假设值
        "preserved_elements": [
            "角色基本信息",
            "案件核心事实",
            "关键时间点"
        ],
        "key_images": full_data['all_images'][:5],  # 前5张图片
        "metadata": {
            "created_date": datetime.now().isoformat(),
            "compression_method": "automatic",
            "version": "v1.0"
        }
    }

    # 保存压缩样本
    output_dir = Path("/home/ubt/桌面/agent-project/data/extracted")
    sample_file = output_dir / "training_sample_compressed.json"

    with open(sample_file, 'w', encoding='utf-8') as f:
        json.dump(training_sample, f, ensure_ascii=False, indent=2)

    print(f"🎯 压缩样本已创建: {sample_file}")
    print(f"   原始长度: {original_length} 字符")
    print(f"   压缩长度: {compressed_length} 字符")
    print(f"   压缩比例: {compression_ratio:.3f}")

    return training_sample


def main():
    """主函数"""
    try:
        # 提取所有文件内容
        extracted_data = process_all_files()

        # 创建压缩样本
        if extracted_data['full_text']:
            create_compressed_sample(extracted_data)

        print(f"\n🚀 数据提取完成！")
        print(f"现在你可以:")
        print(f"1. 查看提取的文本内容")
        print(f"2. 使用压缩样本进行测试")
        print(f"3. 根据需要手动调整质量评分")
        print(f"4. 开始模型训练")

    except Exception as e:
        logger.error(f"处理失败: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()