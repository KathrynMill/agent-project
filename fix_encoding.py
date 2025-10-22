#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
处理DOCX格式的剧本和手册文件
将多个文档内容合并为单个文本文件
使用python-docx确保稳定正确的中文编码处理
"""

import os
import sys
import re

# 配置目录路径
INPUT_DIR = r"c:\Users\11928\Desktop\linshi\1"
OUTPUT_DIR = "output"
SCRIPTS_OUTPUT_FILE = "merged_scripts.txt"
MANUALS_OUTPUT_FILE = "merged_manuals.txt"

# 确保输出目录存在
def ensure_output_dir():
    """确保输出目录存在"""
    output_path = os.path.join(os.getcwd(), OUTPUT_DIR)
    if not os.path.exists(output_path):
        os.makedirs(output_path)
        print(f"✅ 创建输出目录: {output_path}")
    return output_path

# 安装必需的包
def install_required_packages():
    """安装必需的Python包"""
    try:
        # 首先尝试导入，如果失败则安装
        try:
            import docx
            print("✅ python-docx 已安装")
        except ImportError:
            print("🔄 python-docx 未安装，正在安装...")
            import subprocess
            subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
            print("✅ 成功安装 python-docx")
        
        # 尝试docx2txt作为备选
        try:
            import docx2txt
            print("✅ docx2txt 已安装")
        except ImportError:
            print("🔄 docx2txt 未安装，正在安装...")
            import subprocess
            subprocess.check_call([sys.executable, "-m", "pip", "install", "docx2txt"])
            print("✅ 成功安装 docx2txt")
        
        return True
    except Exception as e:
        print(f"❌ 安装依赖包失败: {e}")
        return False

# 使用python-docx提取文本（主要方法）
def extract_text_with_python_docx(file_path):
    """使用python-docx库提取文本"""
    try:
        from docx import Document
        doc = Document(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # 清理文本中的异常字符
                text = clean_text(paragraph.text)
                if text:
                    text_parts.append(text)
        
        # 如果文档有表格，也提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        clean_cell = clean_text(cell.text)
                        if clean_cell:
                            row_text.append(clean_cell)
                if row_text:
                    text_parts.append("\t".join(row_text))
        
        return "\n\n".join(text_parts)
    except Exception as e:
        print(f"🔄 使用python-docx提取失败: {e}")
        return None

# 使用docx2txt提取文本（备选方法）
def extract_text_with_docx2txt(file_path):
    """使用docx2txt库提取文本"""
    try:
        import docx2txt
        text = docx2txt.process(file_path)
        # 清理文本
        return clean_text(text)
    except Exception as e:
        print(f"🔄 使用docx2txt提取失败: {e}")
        return None

# 清理文本
def clean_text(text):
    """清理文本中的异常字符和格式"""
    if not text:
        return ""
    
    # 移除控制字符（除了换行符和制表符）
    text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
    
    # 替换常见的乱码
    text = text.replace('锟斤拷', '"').replace('锟斤拷', '"')
    text = text.replace('缂佸嘲鐢總鍏呮眽', '绷带女人')
    text = text.replace('閺岊垰銇婃径锟�', '雾晓')
    
    # 清理多余的空白字符
    text = re.sub(r'\s+', ' ', text).strip()
    text = re.sub(r'\n\s*\n', '\n\n', text)
    
    return text

# 提取文本（尝试多种方法）
def extract_text_from_docx(file_path):
    """从DOCX文件中提取文本，尝试多种方法"""
    print(f"📄 处理文件: {os.path.basename(file_path)}")
    
    # 优先使用python-docx
    text = extract_text_with_python_docx(file_path)
    if text:
        print(f"✅ 使用python-docx成功提取文本")
        return text
    
    # 如果失败，使用docx2txt作为备选
    text = extract_text_with_docx2txt(file_path)
    if text:
        print(f"✅ 使用docx2txt成功提取文本")
        return text
    
    # 最后尝试直接读取文件（进阶方法）
    try:
        with open(file_path, 'rb') as f:
            binary_data = f.read()
        
        # 尝试多种编码解码
        encodings_to_try = ['utf-8', 'gbk', 'gb2312', 'utf-16']
        for encoding in encodings_to_try:
            try:
                text = binary_data.decode(encoding, errors='replace')
                print(f"🔄 使用{encoding}解码（可能有错误）")
                return clean_text(text)
            except:
                continue
        
        print("❌ 所有方法都失败，无法提取文本")
        return ""
    except Exception as e:
        print(f"❌ 读取文件时出错: {e}")
        return ""

# 处理所有DOCX文件
def process_all_docx_files():
    """处理目录中的所有DOCX文件"""
    if not os.path.exists(INPUT_DIR):
        print(f"❌ 输入目录不存在: {INPUT_DIR}")
        return False
    
    output_path = ensure_output_dir()
    
    # 分离剧本和手册文件
    script_files = []
    manual_files = []
    
    for filename in os.listdir(INPUT_DIR):
        if filename.endswith('.docx') and not filename.startswith('~$'):  # 排除临时文件
            file_path = os.path.join(INPUT_DIR, filename)
            if '手册' in filename:
                manual_files.append((filename, file_path))
            else:
                script_files.append((filename, file_path))
    
    print(f"📋 找到 {len(script_files)} 个剧本文件和 {len(manual_files)} 个手册文件")
    
    # 处理剧本文件
    if script_files:
        script_output = os.path.join(output_path, SCRIPTS_OUTPUT_FILE)
        with open(script_output, 'w', encoding='utf-8') as f:
            for filename, file_path in script_files:
                f.write(f"\n\n===== {filename} =====\n\n")
                text = extract_text_from_docx(file_path)
                if text:
                    f.write(text)
                    f.write("\n\n")
                else:
                    print(f"❌ 无法提取 {filename} 的内容")
        
        # 检查文件大小和内容
        if os.path.exists(script_output):
            file_size = os.path.getsize(script_output)
            print(f"✅ 剧本文件已保存: {script_output} ({file_size:,} 字节)")
    
    # 处理手册文件
    if manual_files:
        manual_output = os.path.join(output_path, MANUALS_OUTPUT_FILE)
        with open(manual_output, 'w', encoding='utf-8') as f:
            for filename, file_path in manual_files:
                f.write(f"\n\n===== {filename} =====\n\n")
                text = extract_text_from_docx(file_path)
                if text:
                    f.write(text)
                    f.write("\n\n")
                else:
                    print(f"❌ 无法提取 {filename} 的内容")
        
        # 检查文件大小和内容
        if os.path.exists(manual_output):
            file_size = os.path.getsize(manual_output)
            print(f"✅ 手册文件已保存: {manual_output} ({file_size:,} 字节)")
    
    return True

# 主函数
def main():
    """主函数"""
    print("🚀 开始处理DOCX文件")
    
    # 安装必需的依赖
    if not install_required_packages():
        print("❌ 无法继续，缺少必需的依赖")
        return 1
    
    # 处理所有文件
    if process_all_docx_files():
        print("🎉 所有文件处理完成！")
        return 0
    else:
        print("❌ 文件处理失败")
        return 1


if __name__ == "__main__":
    sys.exit(main())