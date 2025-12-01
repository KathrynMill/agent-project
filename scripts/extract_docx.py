#!/usr/bin/env python3
"""
提取Word文档剧本内容
处理包含图片的剧本杀文件
"""

import os
import sys
import json
import logging
from pathlib import Path
from typing import Dict, List, Any, Optional
import base64
from datetime import datetime

# 添加项目根目录到Python路径
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Inches
    import zipfile
    from io import BytesIO
    from PIL import Image
except ImportError as e:
    logger.error(f"缺少必要的依赖包: {e}")
    logger.info("请安装: pip install python-docx pillow")
    sys.exit(1)


class DocxExtractor:
    """Word文档提取器"""

    def __init__(self, docx_path: str):
        self.docx_path = Path(docx_path)
        self.document = None
        self.extracted_content = {
            "text": "",
            "images": [],
            "metadata": {}
        }

    def extract_content(self) -> Dict[str, Any]:
        """提取文档内容"""
        try:
            logger.info(f"提取文档: {self.docx_path}")

            # 读取Word文档
            self.document = Document(self.docx_path)

            # 提取文本内容
            self._extract_text()

            # 提取图片信息
            self._extract_images()

            # 提取元数据
            self._extract_metadata()

            return self.extracted_content

        except Exception as e:
            logger.error(f"提取文档失败: {e}")
            raise

    def _extract_text(self):
        """提取文本内容"""
        paragraphs = []

        for paragraph in self.document.paragraphs:
            text = paragraph.text.strip()
            if text:
                paragraphs.append(text)

        self.extracted_content["text"] = "\n".join(paragraphs)
        logger.info(f"提取了 {len(paragraphs)} 个段落")

    def _extract_images(self):
        """提取图片信息"""
        images = []

        try:
            # 从Word文档的zip文件中提取图片
            with zipfile.ZipFile(self.docx_path, 'r') as zip_file:
                # 查找媒体文件夹中的图片
                for file in zip_file.filelist:
                    if file.filename.startswith('word/media/') and \
                       file.filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):

                        # 提取图片数据
                        image_data = zip_file.read(file.filename)

                        # 转换为base64
                        image_base64 = base64.b64encode(image_data).decode('utf-8')

                        # 获取图片信息
                        try:
                            image = Image.open(BytesIO(image_data))
                            width, height = image.size
                            format_name = image.format

                            images.append({
                                "filename": os.path.basename(file.filename),
                                "width": width,
                                "height": height,
                                "format": format_name,
                                "size_bytes": len(image_data),
                                "base64": image_base64,
                                "description": self._generate_image_description(file.filename, image)
                            })
                        except Exception as e:
                            logger.warning(f"处理图片失败 {file.filename}: {e}")
                            images.append({
                                "filename": os.path.basename(file.filename),
                                "size_bytes": len(image_data),
                                "base64": image_base64,
                                "error": str(e)
                            })

            self.extracted_content["images"] = images
            logger.info(f"提取了 {len(images)} 张图片")

        except Exception as e:
            logger.error(f"提取图片失败: {e}")

    def _generate_image_description(self, filename: str, image) -> str:
        """生成图片描述"""
        # 根据文件名生成基础描述
        if "柯太太" in filename:
            return "柯太太的照片/证件照"
        elif "柯少爷" in filename:
            return "柯少爷的照片/证件照"
        elif "云晴" in filename:
            return "云晴的照片/证件照"
        elif "零四" in filename:
            return "零四的照片/证件照"
        elif "雾晓" in filename:
            return "雾晓的照片/证件照"
        elif "手册" in filename or "线索" in filename:
            return f"线索图片/手册页 - {os.path.basename(filename)}"
        elif "时间线" in filename:
            return "时间线/时间轴图片"
        else:
            return f"剧情相关图片 - {os.path.basename(filename)}"

    def _extract_metadata(self):
        """提取文档元数据"""
        try:
            # 基础信息
            self.extracted_content["metadata"] = {
                "filename": self.docx_path.name,
                "file_size": self.docx_path.stat().st_size,
                "extracted_at": datetime.now().isoformat(),
                "paragraph_count": len(self.document.paragraphs),
                "image_count": len(self.extracted_content["images"])
            }

            # 从文件名推断角色
            filename = self.docx_path.name
            if "柯太太" in filename:
                self.extracted_content["role"] = "柯太太"
            elif "柯少爷" in filename:
                self.extracted_content["role"] = "柯少爷"
            elif "云晴" in filename:
                self.extracted_content["role"] = "云晴"
            elif "零四" in filename:
                self.extracted_content["role"] = "零四"
            elif "雾晓" in filename:
                self.extracted_content["role"] = "雾晓"
            elif "手册" in filename:
                self.extracted_content["role"] = "游戏手册"
            elif "线索" in filename:
                self.extracted_content["role"] = "线索材料"
            else:
                self.extracted_content["role"] = "未知"

        except Exception as e:
            logger.error(f"提取元数据失败: {e}")


class ScriptProcessor:
    """剧本处理器"""

    def __init__(self, docx_files: List[str]):
        self.docx_files = docx_files
        self.extracted_data = []
        self.processed_script = {}

    def process_all_files(self) -> Dict[str, Any]:
        """处理所有文件"""
        logger.info(f"开始处理 {len(self.docx_files)} 个文件")

        for docx_file in self.docx_files:
            try:
                extractor = DocxExtractor(docx_file)
                content = extractor.extract_content()
                self.extracted_data.append(content)

            except Exception as e:
                logger.error(f"处理文件失败 {docx_file}: {e}")

        # 组合完整的剧本
        self._combine_script()

        return self.processed_script

    def _combine_script(self):
        """组合完整的剧本"""
        logger.info("组合完整剧本...")

        # 分类提取的内容
        character_scripts = {}
        game_manual = None
        clues = None

        for data in self.extracted_data:
            role = data.get("metadata", {}).get("role", "unknown")

            if role in ["柯太太", "柯少爷", "云晴", "零四", "雾晓"]:
                character_scripts[role] = data
            elif role == "游戏手册":
                game_manual = data
            elif role == "线索材料":
                clues = data
            else:
                logger.warning(f"未知角色类型: {role}")

        # 构建完整剧本
        self.processed_script = {
            "script_id": "柯家庄园谋杀案",
            "title": "柯家庄园谋杀案",
            "metadata": {
                "total_characters": len(character_scripts),
                "total_files": len(self.extracted_data),
                "has_images": any(len(data["images"]) > 0 for data in self.extracted_data),
                "processed_at": datetime.now().isoformat()
            },
            "characters": character_scripts,
            "game_manual": game_manual,
            "clues": clues,
            "full_text": self._build_full_text(character_scripts, game_manual, clues),
            "key_images": self._extract_key_images(self.extracted_data)
        }

    def _build_full_text(self, characters: Dict, manual: Dict, clues: Dict) -> str:
        """构建完整文本"""
        text_parts = []

        # 添加标题和概述
        text_parts.append("剧本标题：柯家庄园谋杀案")
        text_parts.append("\n人物介绍：")

        # 添加角色介绍
        for role, data in characters.items():
            text_parts.append(f"\n{role}：")
            text_parts.append(data["text"][:500] + "..." if len(data["text"]) > 500 else data["text"])

        # 添加游戏手册
        if manual:
            text_parts.append("\n\n游戏手册：")
            text_parts.append(manual["text"])

        # 添加线索
        if clues:
            text_parts.append("\n\n线索材料：")
            text_parts.append(clues["text"])

        return "\n".join(text_parts)

    def _extract_key_images(self, data_list: List[Dict]) -> List[Dict]:
        """提取关键图片"""
        key_images = []

        for data in data_list:
            for img in data["images"]:
                # 跳过错误图片
                if "error" in img:
                    continue

                # 添加上下文信息
                role = data.get("metadata", {}).get("role", "unknown")
                img_info = {
                    "context": f"来自{role}的文件",
                    "description": img.get("description", ""),
                    **{k: v for k, v in img.items() if k not in ["description"]}
                }
                key_images.append(img_info)

        logger.info(f"提取了 {len(key_images)} 张关键图片")
        return key_images

    def save_training_data(self, output_path: str):
        """保存为训练数据格式"""
        logger.info(f"保存训练数据到: {output_path}")

        # 基础训练样本
        training_sample = {
            "script_id": self.processed_script["script_id"],
            "original_script": self.processed_script["full_text"],
            "title": self.processed_script["title"],
            "character_count": len(self.processed_script["characters"]),
            "image_count": len(self.processed_script["key_images"]),
            "metadata": self.processed_script["metadata"],
            "key_images": self.processed_script["key_images"]
        }

        # 保存完整数据
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(training_sample, f, ensure_ascii=False, indent=2)

        # 保存详细提取数据
        detailed_path = output_path.replace('.json', '_detailed.json')
        with open(detailed_path, 'w', encoding='utf-8') as f:
            json.dump(self.processed_script, f, ensure_ascii=False, indent=2)

        logger.info(f"数据已保存到: {output_path}")
        logger.info(f("详细数据已保存到: {detailed_path}")

        return training_sample


def main():
    """主函数"""
    # 获取docx文件列表
    docx_files = [
        "/home/ubt/桌面/agent-project/01 绷带女人 柯太太_QQ浏览器转格式.docx",
        "/home/ubt/桌面/agent-project/02 年轻男子 柯少爷_QQ浏览器转格式.docx",
        "/home/ubt/桌面/agent-project/03 女仆 云晴_QQ浏览器转格式.docx",
        "/home/ubt/桌面/agent-project/04 胡茬男人 零四_QQ浏览器转格式 (1).docx",
        "/home/ubt/桌面/agent-project/05 洋裙女子 雾晓_QQ浏览器转格式.docx",
        "/home/ubt/桌面/agent-project/手册_QQ浏览器转格式.docx",
        "/home/ubt/桌面/agent-project/线索_QQ浏览器转格式.docx"
    ]

    # 过滤存在的文件
    existing_files = [f for f in docx_files if os.path.exists(f)]

    if not existing_files:
        logger.error("没有找到任何docx文件")
        return

    logger.info(f"找到 {len(existing_files)} 个docx文件")

    # 处理文件
    processor = ScriptProcessor(existing_files)
    script_data = processor.process_all_files()

    # 输出统计信息
    print("\n" + "="*50)
    print("剧本提取统计")
    print("="*50)
    print(f"剧本标题: {script_data['title']}")
    print(f"角色数量: {script_data['metadata']['total_characters']}")
    print(f"文件数量: {script_data['metadata']['total_files']}")
    print(f"包含图片: {script_data['metadata']['has_images']}")
    print(f"图片总数: {len(script_data['key_images'])}")
    print(f"文本长度: {len(script_data['full_text'])} 字符")

    print("\n角色列表:")
    for role in script_data["characters"].keys():
        print(f"  - {role}")

    print(f"\n关键图片示例:")
    for i, img in enumerate(script_data["key_images"][:3]):
        print(f"  {i+1}. {img['description']} ({img['width']}x{img['height']})")

    # 保存训练数据
    output_dir = Path("/home/ubt/桌面/agent-project/data/samples")
    output_dir.mkdir(parents=True, exist_ok=True)

    output_file = output_dir / "柯家庄园谋杀案.json"
    training_sample = processor.save_training_data(str(output_file))

    print(f"\n训练数据已保存")
    print(f"文件路径: {output_file}")
    print(f"文本长度: {len(training_sample['original_script'])} 字符")
    print(f"图片数量: {training_sample['image_count']}")


if __name__ == "__main__":
    main()